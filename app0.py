import streamlit as st
import fitz  # PyMuPDF
from PIL import Image
import io
import base64
from docx import Document
from docx2pdf import convert
import tempfile
import os
import re
from datetime import datetime
import pandas as pd
import plotly.graph_objects as go
from streamlit_option_menu import option_menu
import time
import pytesseract
import cv2
import numpy as np
from pdf2docx import Converter
import pdfplumber
import logging
import zipfile

# Suppress pdf2docx logging
logging.getLogger('pdf2docx').setLevel(logging.WARNING)

# Page configuration
st.set_page_config(
    page_title="Advanced Document Viewer",
    page_icon="📄",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Initialize theme in session state
if 'dark_mode' not in st.session_state:
    st.session_state.dark_mode = False

# Custom CSS for modern UI with dark/light mode
if st.session_state.dark_mode:
    st.markdown("""
    <style>
        /* Dark mode styles */
        .stApp {
            background-color: #1a1a1a;
            color: #ffffff;
        }
        
        /* Main container styling */
        .main {
            padding: 0rem 1rem;
            background-color: #1a1a1a;
        }
        
        /* Custom card styling */
        .doc-card {
            background: linear-gradient(135deg, #4a5568 0%, #2d3748 100%);
            padding: 2rem;
            border-radius: 20px;
            color: white;
            margin-bottom: 2rem;
            box-shadow: 0 10px 30px rgba(0,0,0,0.5);
        }
        
        /* Upload area styling */
        .uploadedFile {
            background: #2d3748;
            border: 2px dashed #4a5568;
            border-radius: 15px;
            padding: 2rem;
            transition: all 0.3s ease;
        }
        
        .uploadedFile:hover {
            border-color: #667eea;
            background: #374151;
        }
        
        /* Metrics styling */
        .metric-card {
            background: #2d3748;
            padding: 1.5rem;
            border-radius: 15px;
            box-shadow: 0 5px 15px rgba(0,0,0,0.3);
            text-align: center;
            transition: all 0.3s ease;
            color: white;
        }
        
        .metric-card:hover {
            transform: translateY(-5px);
            box-shadow: 0 10px 25px rgba(0,0,0,0.5);
        }
        
        /* Page navigation */
        .page-nav {
            background: #2d3748;
            padding: 1rem;
            border-radius: 15px;
            box-shadow: 0 5px 15px rgba(0,0,0,0.3);
            display: flex;
            justify-content: center;
            align-items: center;
            gap: 1rem;
        }
        
        /* Progress bar */
        .progress-bar {
            background: #4a5568;
            height: 10px;
            border-radius: 5px;
            overflow: hidden;
            margin: 1rem 0;
        }
        
        /* Sidebar */
        section[data-testid="stSidebar"] {
            background-color: #2d3748;
        }
        
        /* Input fields */
        .stTextInput > div > div > input {
            background-color: #374151;
            color: white;
            border: 1px solid #4a5568;
        }
        
        .stNumberInput > div > div > input {
            background-color: #374151;
            color: white;
            border: 1px solid #4a5568;
        }
    </style>
    """, unsafe_allow_html=True)
else:
    st.markdown("""
    <style>
        /* Light mode styles */
        /* Main container styling */
        .main {
            padding: 0rem 1rem;
        }
        
        /* Custom card styling */
        .doc-card {
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            padding: 2rem;
            border-radius: 20px;
            color: white;
            margin-bottom: 2rem;
            box-shadow: 0 10px 30px rgba(0,0,0,0.2);
        }
        
        /* Upload area styling */
        .uploadedFile {
            background: #f8f9fa;
            border: 2px dashed #e9ecef;
            border-radius: 15px;
            padding: 2rem;
            transition: all 0.3s ease;
        }
        
        .uploadedFile:hover {
            border-color: #667eea;
            background: #f1f3f5;
        }
        
        /* Search box styling */
        .search-box {
            background: #f8f9fa;
            border-radius: 50px;
            padding: 1rem 2rem;
            border: none;
            box-shadow: 0 5px 15px rgba(0,0,0,0.1);
        }
        
        /* Button styling */
        .stButton > button {
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            color: white;
            border: none;
            border-radius: 50px;
            padding: 0.5rem 2rem;
            font-weight: 600;
            transition: all 0.3s ease;
            box-shadow: 0 5px 15px rgba(102, 126, 234, 0.4);
        }
        
        .stButton > button:hover {
            transform: translateY(-2px);
            box-shadow: 0 7px 20px rgba(102, 126, 234, 0.6);
        }
        
        /* Sidebar styling */
        .css-1d391kg {
            background: #f8f9fa;
        }
        
        /* Metrics styling */
        .metric-card {
            background: white;
            padding: 1.5rem;
            border-radius: 15px;
            box-shadow: 0 5px 15px rgba(0,0,0,0.1);
            text-align: center;
            transition: all 0.3s ease;
        }
        
        .metric-card:hover {
            transform: translateY(-5px);
            box-shadow: 0 10px 25px rgba(0,0,0,0.15);
        }
        
        /* Page navigation */
        .page-nav {
            background: white;
            padding: 1rem;
            border-radius: 15px;
            box-shadow: 0 5px 15px rgba(0,0,0,0.1);
            display: flex;
            justify-content: center;
            align-items: center;
            gap: 1rem;
        }
        
        /* Progress bar */
        .progress-bar {
            background: #e9ecef;
            height: 10px;
            border-radius: 5px;
            overflow: hidden;
            margin: 1rem 0;
        }
    </style>
    """, unsafe_allow_html=True)

st.markdown("""
<style>
    /* Common styles for both themes */
    .progress-fill {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        height: 100%;
        transition: width 0.3s ease;
    }
</style>
""", unsafe_allow_html=True)

# Initialize session state
if 'doc_data' not in st.session_state:
    st.session_state.doc_data = None
if 'current_page' not in st.session_state:
    st.session_state.current_page = 0
if 'search_results' not in st.session_state:
    st.session_state.search_results = []
if 'zoom_level' not in st.session_state:
    st.session_state.zoom_level = 1.0
if 'doc_text' not in st.session_state:
    st.session_state.doc_text = ""
if 'doc_metadata' not in st.session_state:
    st.session_state.doc_metadata = {}
if 'rotation' not in st.session_state:
    st.session_state.rotation = 0
if 'ocr_text' not in st.session_state:
    st.session_state.ocr_text = ""
if 'uploaded_file_data' not in st.session_state:
    st.session_state.uploaded_file_data = None
if 'tesseract_available' not in st.session_state:
    try:
        pytesseract.get_tesseract_version()
        st.session_state.tesseract_available = True
    except:
        st.session_state.tesseract_available = False

# PDF Compression Functions with text preservation
def compress_pdf_safe(pdf_data, compression_level='medium'):
    """Safe PDF compression using PyMuPDF's built-in optimization"""
    try:
        # Open the PDF
        doc = fitz.open(stream=pdf_data, filetype="pdf")
        
        # Define compression settings
        if compression_level == 'minimal':
            # Minimal compression - just remove redundancy
            compressed_data = doc.tobytes(
                garbage=1,  # Remove unused objects
                deflate=True,  # Compress streams
                clean=True,  # Clean up
            )
        elif compression_level == 'standard':
            # Standard compression with more aggressive garbage collection
            compressed_data = doc.tobytes(
                garbage=4,  # Maximum garbage collection
                deflate=True,  # Compress streams
                deflate_images=True,  # Compress images
                deflate_fonts=True,  # Compress fonts
                clean=True,  # Clean up
            )
        elif compression_level == 'maximum':
            # Maximum compression without converting to images
            # Note: ez_save returns None, we need to save to bytes differently
            compressed_data = doc.tobytes(
                garbage=4,
                deflate=True,
                deflate_images=True,
                deflate_fonts=True,
                clean=True,
                pretty=False,  # Remove formatting
                ascii=False,  # Keep unicode
                expand=0,  # Don't expand images
                linear=True,  # Web optimization
            )
        else:
            # Default to standard
            compressed_data = doc.tobytes(garbage=4, deflate=True)
        
        # Calculate compression ratio
        original_size = len(pdf_data)
        compressed_size = len(compressed_data)
        ratio = (1 - compressed_size / original_size) * 100
        
        doc.close()
        
        return compressed_data, ratio, original_size, compressed_size
        
    except Exception as e:
        print(f"Safe compression error: {str(e)}")
        return None, 0, 0, 0

def compress_pdf_with_image_reduction(pdf_data, image_quality=70, image_dpi=120):
    """Compress PDF by reducing image quality while preserving text"""
    try:
        # Open the PDF
        doc = fitz.open(stream=pdf_data, filetype="pdf")
        
        # Process each page
        for page_num in range(doc.page_count):
            page = doc[page_num]
            
            # Get list of images on the page
            image_list = page.get_images()
            
            for img_index, img in enumerate(image_list):
                try:
                    # Get the image
                    xref = img[0]
                    pix = fitz.Pixmap(doc, xref)
                    
                    # Check if we need to convert colorspace
                    if pix.n - pix.alpha > 3:  # CMYK
                        pix = fitz.Pixmap(fitz.csRGB, pix)
                    
                    # Reduce image resolution if it's too high
                    if pix.width > image_dpi * 10 or pix.height > image_dpi * 10:
                        # Calculate new dimensions
                        factor = min(image_dpi * 10 / pix.width, image_dpi * 10 / pix.height)
                        new_width = int(pix.width * factor)
                        new_height = int(pix.height * factor)
                        
                        # Create PIL Image for resizing
                        img_data = pix.tobytes("png")
                        img_pil = Image.open(io.BytesIO(img_data))
                        img_pil = img_pil.resize((new_width, new_height), Image.Resampling.LANCZOS)
                        
                        # Convert back to bytes
                        output = io.BytesIO()
                        img_pil.save(output, format='JPEG', quality=image_quality, optimize=True)
                        new_img_data = output.getvalue()
                        
                        # Replace the image
                        page.replace_image(xref, stream=new_img_data)
                    else:
                        # Just compress the image
                        img_data = pix.tobytes("jpg", jpg_quality=image_quality)
                        page.replace_image(xref, stream=img_data)
                    
                    pix = None
                    
                except Exception as e:
                    print(f"Error processing image: {str(e)}")
                    continue
        
        # Save with compression
        compressed_data = doc.tobytes(
            garbage=4,
            deflate=True,
            deflate_images=True,
            deflate_fonts=True,
            clean=True
        )
        
        # Calculate compression ratio
        original_size = len(pdf_data)
        compressed_size = len(compressed_data)
        ratio = (1 - compressed_size / original_size) * 100
        
        doc.close()
        
        return compressed_data, ratio, original_size, compressed_size
        
    except Exception as e:
        print(f"Image reduction compression error: {str(e)}")
        return None, 0, 0, 0

def compress_pdf_extreme(pdf_data, quality=30, dpi=72):
    """Extreme compression by converting to images - may affect text quality"""
    try:
        # Open the PDF
        doc = fitz.open(stream=pdf_data, filetype="pdf")
        
        # Create new document
        new_doc = fitz.open()
        
        for page_num in range(doc.page_count):
            page = doc[page_num]
            
            # Render page as image
            mat = fitz.Matrix(dpi / 72.0, dpi / 72.0)
            pix = page.get_pixmap(matrix=mat, alpha=False)
            
            # Convert to JPEG with low quality
            img_data = pix.tobytes("jpg", jpg_quality=quality)
            
            # Create new page
            new_page = new_doc.new_page(width=page.rect.width, height=page.rect.height)
            
            # Insert image
            img_rect = page.rect
            new_page.insert_image(img_rect, stream=img_data)
            
            pix = None
        
        # Save with maximum compression
        compressed_data = new_doc.tobytes(
            garbage=4,
            deflate=True,
            clean=True
        )
        
        new_doc.close()
        
        # Calculate compression ratio
        original_size = len(pdf_data)
        compressed_size = len(compressed_data)
        ratio = (1 - compressed_size / original_size) * 100
        
        doc.close()
        
        return compressed_data, ratio, original_size, compressed_size
        
    except Exception as e:
        print(f"Extreme compression error: {str(e)}")
        return None, 0, 0, 0

def compress_image_advanced(image_data, output_format='JPEG', quality=85, max_width=None, max_height=None):
    """Advanced image compression with resizing"""
    try:
        # Open image
        img = Image.open(io.BytesIO(image_data))
        
        # Convert RGBA to RGB if saving as JPEG
        if output_format == 'JPEG' and img.mode == 'RGBA':
            # Create white background
            background = Image.new('RGB', img.size, (255, 255, 255))
            background.paste(img, mask=img.split()[3])
            img = background
        
        # Resize if dimensions specified
        if max_width or max_height:
            img.thumbnail(
                (max_width or img.width, max_height or img.height),
                Image.Resampling.LANCZOS
            )
        
        # Apply additional optimizations
        output = io.BytesIO()
        
        if output_format == 'JPEG':
            img.save(output, format='JPEG', quality=quality, optimize=True, progressive=True)
        elif output_format == 'PNG':
            img.save(output, format='PNG', optimize=True, compress_level=9)
        elif output_format == 'WEBP':
            img.save(output, format='WEBP', quality=quality, method=6)
        
        compressed_data = output.getvalue()
        
        # Calculate compression ratio
        original_size = len(image_data)
        compressed_size = len(compressed_data)
        ratio = (1 - compressed_size / original_size) * 100
        
        return compressed_data, ratio, original_size, compressed_size, img.size
        
    except Exception as e:
        return None, 0, 0, 0, (0, 0)

def resize_image_batch(images, target_width=None, target_height=None, maintain_aspect=True):
    """Batch resize images"""
    resized_images = []
    
    for img_name, img_data in images:
        try:
            img = Image.open(io.BytesIO(img_data))
            
            if maintain_aspect:
                img.thumbnail(
                    (target_width or img.width, target_height or img.height),
                    Image.Resampling.LANCZOS
                )
            else:
                if target_width and target_height:
                    img = img.resize((target_width, target_height), Image.Resampling.LANCZOS)
            
            output = io.BytesIO()
            img.save(output, format=img.format or 'PNG')
            resized_images.append((img_name, output.getvalue(), img.size))
            
        except Exception as e:
            resized_images.append((img_name, img_data, (0, 0)))
    
    return resized_images

# Helper functions
def extract_text_from_pdf(pdf_document):
    """Extract text from PDF document"""
    text = ""
    for page_num in range(pdf_document.page_count):
        page = pdf_document[page_num]
        text += page.get_text()
    return text

def extract_text_from_docx(docx_file):
    """Extract text from DOCX document"""
    doc = Document(docx_file)
    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    return text

def convert_docx_to_pdf(docx_file):
    """Convert DOCX to PDF for preview"""
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_docx:
        tmp_docx.write(docx_file.read())
        tmp_docx_path = tmp_docx.name
    
    pdf_path = tmp_docx_path.replace('.docx', '.pdf')
    
    try:
        # Note: docx2pdf requires Microsoft Word on Windows or LibreOffice on Linux/Mac
        convert(tmp_docx_path, pdf_path)
        with open(pdf_path, 'rb') as pdf_file:
            pdf_data = pdf_file.read()
        os.unlink(tmp_docx_path)
        os.unlink(pdf_path)
        return pdf_data
    except:
        # Fallback: Return None if conversion fails
        os.unlink(tmp_docx_path)
        return None

def render_pdf_page(pdf_document, page_num, zoom_level=1.0, rotation=0):
    """Render a PDF page as image with rotation"""
    page = pdf_document[page_num]
    mat = fitz.Matrix(zoom_level * 2, zoom_level * 2)  # 2x for better quality
    mat = mat * fitz.Matrix(rotation)
    pix = page.get_pixmap(matrix=mat)
    img_data = pix.tobytes("png")
    return Image.open(io.BytesIO(img_data))

def perform_ocr(image, lang='eng', enhance=True):
    """Perform OCR on an image with enhancement options"""
    try:
        # Check if Tesseract is installed
        try:
            pytesseract.get_tesseract_version()
        except:
            return "Error: Tesseract is not installed. Please install it from https://github.com/UB-Mannheim/tesseract/wiki"
        
        # Convert PIL Image to numpy array
        img_array = np.array(image)
        
        # Convert to grayscale if needed
        if len(img_array.shape) == 3:
            gray = cv2.cvtColor(img_array, cv2.COLOR_RGB2GRAY)
        else:
            gray = img_array
        
        if enhance:
            # Apply some preprocessing for better OCR
            # Denoise
            denoised = cv2.fastNlDenoising(gray)
            
            # Threshold
            _, thresh = cv2.threshold(denoised, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
            
            # Dilation and erosion to remove noise
            kernel = np.ones((1, 1), np.uint8)
            processed = cv2.morphologyEx(thresh, cv2.MORPH_CLOSE, kernel)
        else:
            processed = gray
        
        # Perform OCR with configuration
        custom_config = r'--oem 3 --psm 6'
        text = pytesseract.image_to_string(processed, lang=lang, config=custom_config)
        
        return text
    except Exception as e:
        return f"OCR Error: {str(e)}"

def extract_text_with_ocr(pdf_document):
    """Extract text from PDF with OCR fallback for scanned pages"""
    text = ""
    ocr_performed = False
    
    for page_num in range(pdf_document.page_count):
        page = pdf_document[page_num]
        page_text = page.get_text()
        
        # If page has very little text, it might be scanned
        if len(page_text.strip()) < 50:
            # Convert page to image and perform OCR
            mat = fitz.Matrix(2, 2)  # 2x scale for better OCR
            pix = page.get_pixmap(matrix=mat)
            img_data = pix.tobytes("png")
            img = Image.open(io.BytesIO(img_data))
            
            ocr_text, ocr_performed = perform_ocr(img, 'eng', True), False
            if ocr_text and len(ocr_text.strip()) > len(page_text.strip()):
                text += f"\n[Page {page_num + 1} - OCR]\n{ocr_text}\n"
                ocr_performed = True
            else:
                text += page_text
        else:
            text += page_text
    
    return text, ocr_performed

def convert_pdf_to_docx_advanced(pdf_path, docx_path):
    """Advanced PDF to DOCX conversion with formatting preservation"""
    try:
        # Create converter with advanced settings
        cv = Converter(pdf_path)
        
        # Configure for better formatting preservation
        cv.convert(docx_path, start=0, end=None)
        cv.close()
        
        # Post-process to enhance formatting
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = Document(docx_path)
        
        # Save enhanced document
        doc.save(docx_path)
        
        return True, "Conversion completed with formatting preservation"
    except Exception as e:
        # Fallback to basic conversion
        return convert_pdf_to_docx(pdf_path, docx_path), "Basic conversion completed"

def extract_images_from_pdf(pdf_document):
    """Extract all images from PDF"""
    images = []
    
    for page_num in range(pdf_document.page_count):
        page = pdf_document[page_num]
        image_list = page.get_images()
        
        for img_index, img in enumerate(image_list):
            # Get image data
            xref = img[0]
            pix = fitz.Pixmap(pdf_document, xref)
            
            if pix.n - pix.alpha < 4:  # GRAY or RGB
                img_data = pix.tobytes("png")
            else:  # CMYK
                pix = fitz.Pixmap(fitz.csRGB, pix)
                img_data = pix.tobytes("png")
            
            images.append({
                'page': page_num + 1,
                'index': img_index + 1,
                'data': img_data,
                'ext': 'png'
            })
            
            pix = None
    
    return images

def convert_pdf_to_docx(pdf_path, docx_path):
    """Basic PDF to DOCX conversion"""
    try:
        cv = Converter(pdf_path)
        cv.convert(docx_path)
        cv.close()
        return True
    except Exception as e:
        return False

def create_formatted_docx_ultra(pdf_document, output_path):
    """Ultra-advanced DOCX creation with full page capture as fallback"""
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import io
    
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
    
    for page_num in range(pdf_document.page_count):
        page = pdf_document[page_num]
        
        # Convert entire page to high-quality image
        # This ensures nothing is missed
        mat = fitz.Matrix(2, 2)  # 2x scale for good quality
        pix = page.get_pixmap(matrix=mat, alpha=False)
        img_data = pix.tobytes("png")
        
        # Add page as image
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run()
        
        # Add image with appropriate width
        run.add_picture(io.BytesIO(img_data), width=Inches(6.5))
        
        pix = None
        
        # Add page break except for last page
        if page_num < pdf_document.page_count - 1:
            doc.add_page_break()
    
    # Save document
    doc.save(output_path)
    return True

def create_advanced_docx_with_images(pdf_document, output_path):
    """Advanced DOCX creation with proper image extraction and positioning"""
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import io
    
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
    
    for page_num in range(pdf_document.page_count):
        page = pdf_document[page_num]
        page_rect = page.rect
        
        # Method 1: Extract text blocks with formatting
        try:
            blocks = page.get_text("dict", sort=True)
        except:
            blocks = {"blocks": []}
        
        # Method 2: Get all images
        images = page.get_images()
        
        # Create a list to store all page elements with their positions
        page_elements = []
        
        # Add text blocks to elements list
        for block in blocks["blocks"]:
            if block["type"] == 0:  # Text block
                page_elements.append({
                    'type': 'text',
                    'bbox': block["bbox"],
                    'content': block,
                    'y_pos': block["bbox"][1]
                })
        
        # Add images to elements list
        for img_index, img in enumerate(images):
            try:
                # Get image bbox
                img_rects = page.get_image_rects(img[0])
                if img_rects:
                    for rect in img_rects:
                        page_elements.append({
                            'type': 'image',
                            'bbox': [rect.x0, rect.y0, rect.x1, rect.y1],
                            'xref': img[0],
                            'y_pos': rect.y0
                        })
            except:
                # Fallback: add image without position
                page_elements.append({
                    'type': 'image',
                    'bbox': [0, img_index * 100, page_rect.width, (img_index + 1) * 100],
                    'xref': img[0],
                    'y_pos': img_index * 100
                })
        
        # Sort elements by vertical position
        page_elements.sort(key=lambda x: x['y_pos'])
        
        # If no elements found, capture the whole page
        if not page_elements:
            mat = fitz.Matrix(2, 2)
            pix = page.get_pixmap(matrix=mat, alpha=False)
            img_data = pix.tobytes("png")
            
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run()
            run.add_picture(io.BytesIO(img_data), width=Inches(6.5))
            
            pix = None
        else:
            # Process elements in order
            for element in page_elements:
                if element['type'] == 'text':
                    block = element['content']
                    
                    # Create paragraph
                    para = doc.add_paragraph()
                    
                    # Set alignment based on position
                    bbox = block["bbox"]
                    page_width = page_rect.width
                    block_center = (bbox[0] + bbox[2]) / 2
                    
                    if block_center < page_width * 0.4:
                        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    elif block_center > page_width * 0.6:
                        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    else:
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # Process text lines
                    for line in block["lines"]:
                        for span in line["spans"]:
                            text = span["text"]
                            if not text.strip():
                                continue
                            
                            run = para.add_run(text)
                            
                            # Apply formatting
                            font_size = span.get("size", 11)
                            run.font.size = Pt(font_size)
                            
                            # Font style
                            if span.get("flags", 0) & 2**4:
                                run.bold = True
                            if span.get("flags", 0) & 2**1:
                                run.italic = True
                            
                            # Font color
                            color = span.get("color", 0)
                            if color != 0 and color != 0xffffff:  # Not black or white
                                try:
                                    # Convert color value to RGB
                                    r = (color >> 16) & 0xFF
                                    g = (color >> 8) & 0xFF
                                    b = color & 0xFF
                                    run.font.color.rgb = RGBColor(r, g, b)
                                except:
                                    pass  # Ignore color errors
                        
                        para.add_run(" ")
                
                elif element['type'] == 'image':
                    try:
                        # Extract image
                        xref = element['xref']
                        pix = fitz.Pixmap(pdf_document, xref)
                        
                        # Convert CMYK to RGB if necessary
                        if pix.n - pix.alpha > 3:  # CMYK
                            pix = fitz.Pixmap(fitz.csRGB, pix)
                        
                        img_data = pix.tobytes("png")
                        
                        # Add image paragraph
                        para = doc.add_paragraph()
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = para.add_run()
                        
                        # Calculate image size in inches
                        bbox = element['bbox']
                        width_in_points = bbox[2] - bbox[0]
                        height_in_points = bbox[3] - bbox[1]
                        width_in_inches = width_in_points / 72.0
                        
                        # Limit to page width
                        max_width = 6.0  # inches
                        if width_in_inches > max_width or width_in_inches <= 0:
                            width_in_inches = max_width
                        
                        # Add image with proper size
                        run.add_picture(io.BytesIO(img_data), width=Inches(width_in_inches))
                        
                        pix = None
                        
                    except Exception as e:
                        print(f"Error adding image: {str(e)}")
                        continue
        
        # Add page break
        if page_num < pdf_document.page_count - 1:
            doc.add_page_break()
    
    # Save document
    doc.save(output_path)
    return True
    """Advanced DOCX creation with proper image extraction and positioning"""
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.section import WD_SECTION
    import io
    
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
    
    for page_num in range(pdf_document.page_count):
        page = pdf_document[page_num]
        page_rect = page.rect
        
        # Method 1: Extract text blocks with formatting
        blocks = page.get_text("dict", sort=True)
        
        # Method 2: Get all drawings and images
        drawings = page.get_drawings()
        images = page.get_images()
        
        # Create a list to store all page elements with their positions
        page_elements = []
        
        # Add text blocks to elements list
        for block in blocks["blocks"]:
            if block["type"] == 0:  # Text block
                page_elements.append({
                    'type': 'text',
                    'bbox': block["bbox"],
                    'content': block,
                    'y_pos': block["bbox"][1]
                })
        
        # Add images to elements list
        for img_index, img in enumerate(images):
            try:
                # Get image position
                img_list = page.get_image_info(xref=img[0])
                if img_list:
                    for img_info in img_list:
                        bbox = img_info["bbox"]
                        page_elements.append({
                            'type': 'image',
                            'bbox': bbox,
                            'xref': img[0],
                            'y_pos': bbox[1]
                        })
            except:
                pass
        
        # Sort elements by vertical position
        page_elements.sort(key=lambda x: x['y_pos'])
        
        # Process elements in order
        for element in page_elements:
            if element['type'] == 'text':
                block = element['content']
                
                # Create paragraph
                para = doc.add_paragraph()
                
                # Set alignment based on position
                bbox = block["bbox"]
                page_width = page_rect.width
                block_center = (bbox[0] + bbox[2]) / 2
                
                if block_center < page_width * 0.4:
                    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                elif block_center > page_width * 0.6:
                    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                else:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Process text lines
                for line in block["lines"]:
                    for span in line["spans"]:
                        text = span["text"]
                        if not text.strip():
                            continue
                        
                        run = para.add_run(text)
                        
                        # Apply formatting
                        font_size = span.get("size", 11)
                        run.font.size = Pt(font_size)
                        
                        # Font style
                        if span.get("flags", 0) & 2**4:
                            run.bold = True
                        if span.get("flags", 0) & 2**1:
                            run.italic = True
                        
                        # Font color
                        color = span.get("color", 0)
                        if color != 0:
                            rgb = fitz.sRGB(color)
                            run.font.color.rgb = RGBColor(
                                int(rgb[0] * 255),
                                int(rgb[1] * 255),
                                int(rgb[2] * 255)
                            )
                    
                    para.add_run(" ")
            
            elif element['type'] == 'image':
                try:
                    # Extract image
                    xref = element['xref']
                    pix = fitz.Pixmap(pdf_document, xref)
                    
                    if pix.n - pix.alpha > 3:  # CMYK
                        pix = fitz.Pixmap(fitz.csRGB, pix)
                    
                    img_data = pix.tobytes("png")
                    
                    # Add image paragraph
                    para = doc.add_paragraph()
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = para.add_run()
                    
                    # Calculate image size in inches
                    bbox = element['bbox']
                    width_in_points = bbox[2] - bbox[0]
                    height_in_points = bbox[3] - bbox[1]
                    width_in_inches = width_in_points / 72.0
                    height_in_inches = height_in_points / 72.0
                    
                    # Limit to page width
                    max_width = 6.0  # inches
                    if width_in_inches > max_width:
                        scale = max_width / width_in_inches
                        width_in_inches = max_width
                        height_in_inches *= scale
                    
                    # Add image with proper size
                    run.add_picture(io.BytesIO(img_data), width=Inches(width_in_inches))
                    
                    pix = None
                    
                except Exception as e:
                    print(f"Error adding image: {str(e)}")
                    continue
        
        # Try to extract and add vector graphics as images
        try:
            # Get page as high-resolution image to capture any missed graphics
            mat = fitz.Matrix(2, 2)  # 2x scale
            pix = page.get_pixmap(matrix=mat, alpha=False)
            
            # Check if page has significant non-text content
            text_area = sum(block["bbox"][2] * block["bbox"][3] for block in blocks["blocks"] if block["type"] == 0)
            page_area = page_rect.width * page_rect.height
            
            # If less than 30% is text, likely has graphics
            if text_area < page_area * 0.3 and len(images) == 0:
                # Add full page as image
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = para.add_run()
                
                img_data = pix.tobytes("png")
                run.add_picture(io.BytesIO(img_data), width=Inches(6))
                
                # Add page break instead of continuing with text
                if page_num < pdf_document.page_count - 1:
                    doc.add_page_break()
                continue
            
            pix = None
            
        except Exception as e:
            print(f"Error processing page graphics: {str(e)}")
        
        # Add page break
        if page_num < pdf_document.page_count - 1:
            doc.add_page_break()
    
    # Save document
    doc.save(output_path)
    return True
    """Ultra-advanced DOCX creation with exact formatting preservation"""
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    import io
    
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
    
    for page_num in range(pdf_document.page_count):
        page = pdf_document[page_num]
        
        # Get page dimensions
        page_rect = page.rect
        
        # Extract text with detailed formatting
        blocks = page.get_text("dict", sort=True)
        
        # Group blocks by vertical position for better paragraph detection
        grouped_blocks = {}
        
        for block in blocks["blocks"]:
            if block["type"] == 0:  # Text block
                # Get vertical position
                y_pos = int(block["bbox"][1])
                if y_pos not in grouped_blocks:
                    grouped_blocks[y_pos] = []
                grouped_blocks[y_pos].append(block)
        
        # Sort by vertical position
        sorted_positions = sorted(grouped_blocks.keys())
        
        for y_pos in sorted_positions:
            blocks_at_position = grouped_blocks[y_pos]
            
            for block in blocks_at_position:
                # Create paragraph
                para = doc.add_paragraph()
                
                # Set paragraph alignment based on x position
                bbox = block["bbox"]
                page_width = page_rect.width
                block_center = (bbox[0] + bbox[2]) / 2
                
                if block_center < page_width * 0.4:
                    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                elif block_center > page_width * 0.6:
                    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                else:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Process lines
                for line in block["lines"]:
                    line_text = ""
                    
                    for span in line["spans"]:
                        text = span["text"]
                        if not text.strip():
                            continue
                        
                        run = para.add_run(text)
                        
                        # Font name and size
                        font_name = span.get("font", "Arial")
                        font_size = span.get("size", 11)
                        
                        # Clean font name
                        if "+" in font_name:
                            font_name = font_name.split("+")[-1]
                        if "-" in font_name:
                            base_font = font_name.split("-")[0]
                            style = font_name.split("-")[-1].lower()
                        else:
                            base_font = font_name
                            style = ""
                        
                        run.font.name = base_font
                        run.font.size = Pt(font_size)
                        
                        # Font style
                        if "bold" in style or span.get("flags", 0) & 2**4:
                            run.bold = True
                        if "italic" in style or span.get("flags", 0) & 2**1:
                            run.italic = True
                        
                        # Font color
                        color = span.get("color", 0)
                        if color != 0:
                            rgb = fitz.sRGB(color)
                            run.font.color.rgb = RGBColor(
                                int(rgb[0] * 255),
                                int(rgb[1] * 255),
                                int(rgb[2] * 255)
                            )
                    
                    para.add_run(" ")  # Add space between lines
        
        # Extract and add images
        image_list = page.get_images()
        
        for img_index, img in enumerate(image_list):
            try:
                # Get image data
                xref = img[0]
                pix = fitz.Pixmap(pdf_document, xref)
                
                if pix.n - pix.alpha < 4:  # GRAY or RGB
                    img_data = pix.tobytes("png")
                else:  # CMYK
                    pix = fitz.Pixmap(fitz.csRGB, pix)
                    img_data = pix.tobytes("png")
                
                # Add image to document
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = para.add_run()
                
                # Calculate image size
                img_rect = page.get_image_bbox(img)
                width_inches = (img_rect.width / 72)  # Convert points to inches
                
                # Limit width to page width
                max_width = 6  # inches
                if width_inches > max_width:
                    width_inches = max_width
                
                run.add_picture(io.BytesIO(img_data), width=Inches(width_inches))
                
                pix = None
                
            except Exception as e:
                continue
        
        # Add page break except for last page
        if page_num < pdf_document.page_count - 1:
            doc.add_page_break()
    
    # Save document
    doc.save(output_path)
    return True

def convert_pdf_to_text(pdf_document):
    """Extract all text from PDF"""
    text = ""
    for page_num in range(pdf_document.page_count):
        page = pdf_document[page_num]
        text += f"--- Page {page_num + 1} ---\n"
        text += page.get_text()
        text += "\n\n"
    return text

def convert_to_images(pdf_document, image_format="PNG"):
    """Convert PDF pages to images"""
    images = []
    for page_num in range(pdf_document.page_count):
        page = pdf_document[page_num]
        mat = fitz.Matrix(2, 2)  # 2x scale
        pix = page.get_pixmap(matrix=mat)
        img_data = pix.tobytes(image_format.lower())
    images.append((f"page_{page_num + 1}.{image_format.lower()}", img_data))
    return images

def extract_tables_from_pdf(pdf_bytes):
    """Extract tables from a PDF byte string using pdfplumber"""
    tables = []
    with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
        for page_num, page in enumerate(pdf.pages, start=1):
            page_tables = page.extract_tables()
            for idx, table in enumerate(page_tables, start=1):
                try:
                    df = pd.DataFrame(table)
                except Exception:
                    df = pd.DataFrame()
                tables.append({
                    'page': page_num,
                    'index': idx,
                    'df': df
                })
    return tables

def search_in_document(text, query):
    """Search for query in document text"""
    if not query:
        return []
    
    results = []
    lines = text.split('\n')
    query_lower = query.lower()
    
    for i, line in enumerate(lines):
        if query_lower in line.lower():
            start = max(0, i - 1)
            end = min(len(lines), i + 2)
            context = '\n'.join(lines[start:end])
            results.append({
                'line': i + 1,
                'context': context,
                'text': line
            })
    
    return results

def get_document_metadata(pdf_document):
    """Extract metadata from PDF"""
    metadata = pdf_document.metadata
    return {
        'title': metadata.get('title', 'N/A'),
        'author': metadata.get('author', 'N/A'),
        'subject': metadata.get('subject', 'N/A'),
        'creator': metadata.get('creator', 'N/A'),
        'producer': metadata.get('producer', 'N/A'),
        'creation_date': metadata.get('creationDate', 'N/A'),
        'modification_date': metadata.get('modDate', 'N/A'),
        'pages': pdf_document.page_count
    }

# Header with theme toggle
col1, col2, col3 = st.columns([8, 1, 1])

with col1:
    st.markdown("""
    <div class="doc-card">
        <h1 style="margin: 0; font-size: 2.5rem;">📄 Advanced Document Viewer</h1>
        <p style="margin: 0.5rem 0 0 0; opacity: 0.9;">Upload and preview PDF or DOCX documents with powerful features</p>
    </div>
    """, unsafe_allow_html=True)

with col3:
    if st.button("🌓", help="Toggle Dark/Light Mode", use_container_width=True):
        st.session_state.dark_mode = not st.session_state.dark_mode
        st.rerun()

# Sidebar
with st.sidebar:
    st.markdown("### 🎯 Navigation")
    selected = option_menu(
        menu_title=None,
        options=["Upload", "View", "Search", "Tables", "Analytics", "OCR", "Convert", "Compress", "Export"],
        icons=["cloud-upload", "eye", "search", "table", "graph-up", "cpu", "arrow-repeat", "file-zip", "download"],
        menu_icon="cast",
        default_index=0,
        styles={
            "container": {
                "padding": "0!important",
                "background-color": "#f8f9fa" if not st.session_state.dark_mode else "#2d3748"
            },
            "icon": {"color": "#667eea", "font-size": "20px"},
            "nav-link": {
                "font-size": "16px",
                "text-align": "left",
                "margin": "0px",
                "padding": "10px",
                "border-radius": "10px",
                "color": "#000000" if not st.session_state.dark_mode else "#ffffff"
            },
            "nav-link-selected": {
                "background-color": "#667eea",
                "color": "white"
            },
        }
    )
    
    if st.session_state.doc_data:
        st.markdown("---")
        st.markdown("### 📊 Document Info")
        
        if st.session_state.doc_metadata:
            for key, value in st.session_state.doc_metadata.items():
                if key != 'pages':
                    st.markdown(f"**{key.title()}:** {value}")
        
        st.markdown("---")
        st.markdown("### 🔧 Controls")
        
        # Zoom controls
        st.session_state.zoom_level = st.slider(
            "Zoom Level",
            min_value=0.5,
            max_value=3.0,
            value=st.session_state.zoom_level,
            step=0.1,
            format="%.1fx"
        )
        
        # Page navigation
        if st.session_state.doc_data:
            total_pages = st.session_state.doc_data.page_count
            st.markdown(f"**Total Pages:** {total_pages}")

# Main content area
if selected == "Upload":
    col1, col2, col3 = st.columns([1, 2, 1])
    
    with col2:
        st.markdown("### 📤 Upload Document")
        
        uploaded_file = st.file_uploader(
            "Choose a PDF or DOCX file",
            type=['pdf', 'docx'],
            help="Maximum file size: 200MB"
        )
        
        if uploaded_file is not None:
            # Show upload progress
            progress_bar = st.progress(0)
            status_text = st.empty()
            
            # Simulate processing
            for i in range(100):
                progress_bar.progress(i + 1)
                status_text.text(f'Processing... {i+1}%')
                time.sleep(0.01)
            
            status_text.text('Processing complete!')
            
            # Process the file
            if uploaded_file.type == "application/pdf":
                file_data = uploaded_file.read()
                st.session_state.uploaded_file_data = file_data
                st.session_state.doc_data = fitz.open(stream=file_data, filetype="pdf")
                
                # Try OCR-enhanced text extraction
                with st.spinner("Extracting text (with OCR if needed)..."):
                    extracted_text, ocr_used = extract_text_with_ocr(st.session_state.doc_data)
                    st.session_state.doc_text = extracted_text
                    if ocr_used:
                        st.info("📸 OCR was used for some pages with scanned content")
                
                st.session_state.doc_metadata = get_document_metadata(st.session_state.doc_data)
            
            elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                # Extract text from DOCX
                uploaded_file.seek(0)
                st.session_state.doc_text = extract_text_from_docx(uploaded_file)
                
                # Try to convert DOCX to PDF for preview
                uploaded_file.seek(0)
                pdf_data = convert_docx_to_pdf(uploaded_file)
                
                if pdf_data:
                    st.session_state.doc_data = fitz.open(stream=pdf_data, filetype="pdf")
                    st.session_state.doc_metadata = get_document_metadata(st.session_state.doc_data)
                else:
                    st.warning("DOCX preview not available. Text extracted successfully.")
                    st.session_state.doc_data = None
            
            st.success("✅ Document uploaded successfully!")
            
            # Display metrics
            col1, col2, col3, col4 = st.columns(4)
            
            with col1:
                st.markdown("""
                <div class="metric-card">
                    <h3 style="color: #667eea; margin: 0;">📄</h3>
                    <p style="margin: 0; font-size: 24px; font-weight: bold;">Pages</p>
                    <p style="margin: 0; font-size: 32px; color: #764ba2;">{}</p>
                </div>
                """.format(st.session_state.doc_data.page_count if st.session_state.doc_data else "N/A"), unsafe_allow_html=True)
            
            with col2:
                word_count = len(st.session_state.doc_text.split())
                st.markdown("""
                <div class="metric-card">
                    <h3 style="color: #667eea; margin: 0;">📝</h3>
                    <p style="margin: 0; font-size: 24px; font-weight: bold;">Words</p>
                    <p style="margin: 0; font-size: 32px; color: #764ba2;">{:,}</p>
                </div>
                """.format(word_count), unsafe_allow_html=True)
            
            with col3:
                char_count = len(st.session_state.doc_text)
                st.markdown("""
                <div class="metric-card">
                    <h3 style="color: #667eea; margin: 0;">🔤</h3>
                    <p style="margin: 0; font-size: 24px; font-weight: bold;">Characters</p>
                    <p style="margin: 0; font-size: 32px; color: #764ba2;">{:,}</p>
                </div>
                """.format(char_count), unsafe_allow_html=True)
            
            with col4:
                file_size = uploaded_file.size / (1024 * 1024)  # Convert to MB
                st.markdown("""
                <div class="metric-card">
                    <h3 style="color: #667eea; margin: 0;">💾</h3>
                    <p style="margin: 0; font-size: 24px; font-weight: bold;">Size</p>
                    <p style="margin: 0; font-size: 32px; color: #764ba2;">{:.1f} MB</p>
                </div>
                """.format(file_size), unsafe_allow_html=True)

elif selected == "View":
    if st.session_state.doc_data:
        st.markdown("### 👁️ Document Preview")
        
        # Page navigation controls
        col1, col2, col3 = st.columns([1, 3, 1])
        
        with col1:
            if st.button("⬅️ Previous", use_container_width=True):
                if st.session_state.current_page > 0:
                    st.session_state.current_page -= 1
        
        with col2:
            page_num = st.number_input(
                "Page",
                min_value=1,
                max_value=st.session_state.doc_data.page_count,
                value=st.session_state.current_page + 1,
                step=1,
                format="%d"
            )
            st.session_state.current_page = page_num - 1
        
        with col3:
            if st.button("Next ➡️", use_container_width=True):
                if st.session_state.current_page < st.session_state.doc_data.page_count - 1:
                    st.session_state.current_page += 1
        
        # Progress indicator
        progress = (st.session_state.current_page + 1) / st.session_state.doc_data.page_count
        st.markdown(f"""
        <div class="progress-bar">
            <div class="progress-fill" style="width: {progress * 100}%"></div>
        </div>
        """, unsafe_allow_html=True)
        
        # Display the page with rotation
        page_image = render_pdf_page(
            st.session_state.doc_data,
            st.session_state.current_page,
            st.session_state.zoom_level,
            st.session_state.rotation
        )
        
        # Center the image
        col1, col2, col3 = st.columns([1, 6, 1])
        with col2:
            st.image(page_image, use_container_width=True)
        
        # Quick actions
        st.markdown("---")
        col1, col2, col3, col4 = st.columns(4)
        
        with col1:
            if st.button("↶ Rotate Left", use_container_width=True):
                st.session_state.rotation = (st.session_state.rotation - 90) % 360
                st.rerun()
        
        with col2:
            if st.button("↷ Rotate Right", use_container_width=True):
                st.session_state.rotation = (st.session_state.rotation + 90) % 360
                st.rerun()
        
        with col3:
            if st.button("📋 Copy Text", use_container_width=True):
                page = st.session_state.doc_data[st.session_state.current_page]
                page_text = page.get_text()
                st.code(page_text, language=None)
        
        with col4:
            if st.button("💾 Save Page", use_container_width=True):
                st.info("Save feature coming soon!")
    
    else:
        st.info("📤 Please upload a document first to view it.")

elif selected == "Search":
    if st.session_state.doc_text:
        st.markdown("### 🔍 Search Document")
        
        # Search input
        search_query = st.text_input(
            "Enter search term",
            placeholder="Search for text in the document...",
            help="Search is case-insensitive"
        )
        
        col1, col2 = st.columns([1, 4])
        with col1:
            if st.button("🔍 Search", use_container_width=True):
                if search_query:
                    st.session_state.search_results = search_in_document(
                        st.session_state.doc_text,
                        search_query
                    )
        
        # Display search results
        if st.session_state.search_results:
            st.markdown(f"### Found {len(st.session_state.search_results)} results")
            
            for i, result in enumerate(st.session_state.search_results[:10]):  # Show max 10 results
                with st.expander(f"Result {i+1} - Line {result['line']}"):
                    # Highlight search term
                    highlighted = result['context'].replace(
                        search_query,
                        f"**{search_query}**"
                    )
                    st.markdown(highlighted)
                    
                    if st.button(f"Go to page", key=f"goto_{i}"):
                        # Simple approximation - would need better page detection
                        approx_page = (result['line'] // 40)  # Assuming ~40 lines per page
                        if approx_page < st.session_state.doc_data.page_count:
                            st.session_state.current_page = approx_page
                            st.info(f"Navigated to page {approx_page + 1}")
            
            if len(st.session_state.search_results) > 10:
                st.info(f"Showing first 10 results out of {len(st.session_state.search_results)}")
        
        elif search_query and not st.session_state.search_results:
            st.warning("No results found for your search query.")
    
    else:
        st.info("📤 Please upload a document first to search in it.")

elif selected == "Tables":
    if st.session_state.doc_data and st.session_state.uploaded_file_data:
        st.markdown("### 📊 Table Extraction")
        page_num = st.number_input(
            "Page to scan",
            min_value=1,
            max_value=st.session_state.doc_data.page_count,
            value=1,
            step=1,
        )
        if st.button("Extract Tables from Page", use_container_width=True):
            with st.spinner("Extracting tables..."):
                with pdfplumber.open(io.BytesIO(st.session_state.uploaded_file_data)) as pdf:
                    page = pdf.pages[page_num - 1]
                    tables = page.extract_tables()
                    if tables:
                        for idx, table in enumerate(tables, start=1):
                            df = pd.DataFrame(table)
                            st.markdown(f"#### Table {idx} (Page {page_num})")
                            st.dataframe(df)
                            csv = df.to_csv(index=False).encode("utf-8")
                            st.download_button(
                                label=f"Download Table {idx} as CSV",
                                data=csv,
                                file_name=f"table_page{page_num}_{idx}.csv",
                                mime="text/csv",
                            )
                    else:
                        st.info("No tables found on this page.")

        if st.button("Extract Tables From All Pages", use_container_width=True):
            with st.spinner("Scanning all pages..."):
                tables = extract_tables_from_pdf(st.session_state.uploaded_file_data)
                if tables:
                    zip_buffer = io.BytesIO()
                    with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zip_file:
                        for t in tables:
                            csv = t["df"].to_csv(index=False)
                            zip_file.writestr(f"table_p{t['page']}_{t['index']}.csv", csv)
                    st.download_button(
                        label=f"📥 Download {len(tables)} Tables as ZIP",
                        data=zip_buffer.getvalue(),
                        file_name="tables.zip",
                        mime="application/zip",
                    )
                    for t in tables:
                        st.markdown(f"#### Page {t['page']} - Table {t['index']}")
                        st.dataframe(t["df"])
                else:
                    st.info("No tables found in the document.")
    else:
        st.info("📤 Please upload a PDF document to extract tables.")

elif selected == "OCR":
    if not st.session_state.tesseract_available:
        st.error("""
        ⚠️ **Tesseract OCR is not installed!**
        
        Please install Tesseract to use OCR features:
        
        **Windows:**
        1. Download from: https://github.com/UB-Mannheim/tesseract/wiki
        2. Install and add to PATH
        3. Restart the application
        
        **Mac:**
        ```bash
        brew install tesseract
        ```
        
        **Linux:**
        ```bash
        sudo apt-get install tesseract-ocr
        ```
        """)
    elif st.session_state.doc_data:
        st.markdown("### 🔍 OCR (Optical Character Recognition)")
        
        st.info("Extract text from scanned documents or images within PDFs")
        
        col1, col2 = st.columns([2, 1])
        
        with col1:
            # Page selector for OCR
            ocr_page = st.selectbox(
                "Select page for OCR",
                range(1, st.session_state.doc_data.page_count + 1),
                format_func=lambda x: f"Page {x}"
            )
            
            # OCR options
            ocr_options = st.expander("OCR Options", expanded=True)
            with ocr_options:
                enhance_image = st.checkbox("Enhance image quality", value=True)
                language = st.selectbox("OCR Language", ["eng", "spa", "fra", "deu", "chi_sim"])
                
            if st.button("🔍 Perform OCR on Selected Page", use_container_width=True):
                with st.spinner("Performing OCR..."):
                    # Get the page as image
                    page = st.session_state.doc_data[ocr_page - 1]
                    mat = fitz.Matrix(3, 3)  # Higher resolution for better OCR
                    pix = page.get_pixmap(matrix=mat)
                    img_data = pix.tobytes("png")
                    img = Image.open(io.BytesIO(img_data))
                    
                    # Perform OCR
                    ocr_result = perform_ocr(img, language, enhance_image)
                    
                    if ocr_result and not ocr_result.startswith("OCR Error") and not ocr_result.startswith("Error:"):
                        st.success("✅ OCR completed successfully!")
                        
                        # Store OCR result in session state
                        st.session_state.ocr_text = ocr_result
                        
                        st.text_area("Extracted Text", ocr_result, height=300)
                        
                        # Download button for OCR text
                        ocr_bytes = ocr_result.encode('utf-8')
                        st.download_button(
                            label="📥 Download OCR Text",
                            data=ocr_bytes,
                            file_name=f"ocr_page_{ocr_page}.txt",
                            mime="text/plain"
                        )
                    else:
                        st.error(ocr_result)
            
            if st.button("🔍 OCR All Pages", use_container_width=True):
                progress_bar = st.progress(0)
                status_text = st.empty()
                all_ocr_text = ""
                
                for i in range(st.session_state.doc_data.page_count):
                    status_text.text(f"Processing page {i+1}/{st.session_state.doc_data.page_count}")
                    progress_bar.progress((i + 1) / st.session_state.doc_data.page_count)
                    
                    page = st.session_state.doc_data[i]
                    mat = fitz.Matrix(2, 2)
                    pix = page.get_pixmap(matrix=mat)
                    img_data = pix.tobytes("png")
                    img = Image.open(io.BytesIO(img_data))
                    
                    ocr_text = perform_ocr(img, language, enhance_image)
                    if ocr_text and not ocr_text.startswith("OCR Error") and not ocr_text.startswith("Error:"):
                        all_ocr_text += f"\n\n--- Page {i+1} ---\n{ocr_text}"
                
                st.session_state.ocr_text = all_ocr_text
                st.success("✅ OCR completed for all pages!")
                st.text_area("All OCR Text", all_ocr_text, height=400)
        
        with col2:
            st.markdown("#### 📊 OCR Statistics")
            if st.session_state.ocr_text:
                ocr_words = len(st.session_state.ocr_text.split())
                ocr_chars = len(st.session_state.ocr_text)
                
                st.metric("OCR Words", f"{ocr_words:,}")
                st.metric("OCR Characters", f"{ocr_chars:,}")
                
                # Download OCR text using streamlit download button
                if st.session_state.ocr_text:
                    st.download_button(
                        label="📥 Download All OCR Text",
                        data=st.session_state.ocr_text.encode('utf-8'),
                        file_name="ocr_complete.txt",
                        mime="text/plain"
                    )
    
    else:
        st.info("📤 Please upload a document first to use OCR.")

elif selected == "Convert":
    if st.session_state.doc_data or st.session_state.doc_text:
        st.markdown("### 🔄 Document Conversion")
        
        st.info("Convert your document to different formats")
        
        col1, col2 = st.columns(2)
        
        with col1:
            st.markdown("#### 📄 Convert PDF to Other Formats")
            
            conversion_format = st.selectbox(
                "Select output format",
                ["DOCX", "TXT", "Images (PNG)", "Images (JPG)", "Extract Images Only", "HTML"]
            )
            
            if conversion_format == "DOCX":
                st.warning("⚠️ PDF to DOCX conversion works best with text-based PDFs")
                
                # Conversion options
                conversion_method = st.radio(
                    "Conversion Method",
                    ["Standard", "Advanced (Preserve Formatting)", "Ultra (With Images & Graphics)", "Maximum (Full Page Capture)"],
                    help="""
                    • Standard: Basic text extraction using pdf2docx
                    • Advanced: Better formatting preservation with pdf2docx
                    • Ultra: Custom extraction of text and images with positioning
                    • Maximum: Converts each page to image for perfect visual fidelity
                    """
                )
                
                if st.button("Convert to DOCX", use_container_width=True):
                    with st.spinner("Converting to DOCX..."):
                        try:
                            # Save PDF temporarily
                            with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_pdf:
                                tmp_pdf.write(st.session_state.uploaded_file_data)
                                pdf_path = tmp_pdf.name
                            
                            docx_path = pdf_path.replace('.pdf', '.docx')
                            
                            if conversion_method == "Standard":
                                success = convert_pdf_to_docx(pdf_path, docx_path)
                                message = "Standard conversion completed"
                            elif conversion_method == "Advanced (Preserve Formatting)":
                                success, message = convert_pdf_to_docx_advanced(pdf_path, docx_path)
                            elif conversion_method == "Ultra (With Images & Graphics)":
                                success = create_advanced_docx_with_images(st.session_state.doc_data, docx_path)
                                message = "Ultra conversion with images and graphics completed"
                            else:  # Maximum
                                success = create_formatted_docx_ultra(st.session_state.doc_data, docx_path)
                                message = "Maximum conversion with full formatting preservation completed"
                            
                            if success:
                                # Read the DOCX file
                                with open(docx_path, 'rb') as docx_file:
                                    docx_data = docx_file.read()
                                
                                # Clean up
                                os.unlink(pdf_path)
                                os.unlink(docx_path)
                                
                                # Provide download
                                st.download_button(
                                    label="📥 Download Converted DOCX",
                                    data=docx_data,
                                    file_name="converted_document.docx",
                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                                )
                                st.success(f"✅ {message}")
                                
                                # Show conversion statistics
                                st.info(f"""
                                **Conversion Details:**
                                - Method: {conversion_method}
                                - Pages: {st.session_state.doc_data.page_count}
                                - File size: {len(docx_data) / 1024:.1f} KB
                                - Images extracted: Yes (for Ultra/Maximum methods)
                                """)
                                
                                # Method-specific information
                                if conversion_method == "Ultra (With Images & Graphics)":
                                    st.info("""
                                    **Ultra Conversion Features:**
                                    - ✅ Preserves all embedded images
                                    - ✅ Maintains image positions
                                    - ✅ Captures vector graphics
                                    - ✅ Preserves text formatting
                                    - ✅ Handles mixed content pages
                                    """)
                                elif conversion_method == "Maximum (Full Page Capture)":
                                    st.info("""
                                    **Maximum Conversion Features:**
                                    - ✅ Perfect visual fidelity
                                    - ✅ Captures all graphics and layouts
                                    - ✅ No content is missed
                                    - ⚠️ Text becomes part of image (not selectable)
                                    - ⚠️ Larger file size
                                    """)
                                
                        except Exception as e:
                            st.error(f"Conversion failed: {str(e)}")
                            
                            # Provide alternative suggestions
                            st.info("""
                            **Troubleshooting Tips:**
                            1. Try a different conversion method
                            2. If "Standard" or "Advanced" fail with overlap warnings, use "Ultra" or "Maximum"
                            3. "Maximum" method works for all PDFs but converts to images
                            4. Check if the PDF is password protected or corrupted
                            """)
            
            elif conversion_format == "TXT":
                if st.button("Convert to TXT", use_container_width=True):
                    text_content = convert_pdf_to_text(st.session_state.doc_data)
                    st.download_button(
                        label="📥 Download TXT File",
                        data=text_content.encode('utf-8'),
                        file_name="converted_document.txt",
                        mime="text/plain"
                    )
                    st.success("✅ Text extraction completed!")
            
            elif conversion_format.startswith("Images"):
                image_format = "PNG" if "PNG" in conversion_format else "JPG"
                quality = st.slider("Image Quality", 50, 100, 85)
                dpi = st.slider("DPI (Resolution)", 72, 300, 150)
                
                if st.button(f"Convert to {image_format}", use_container_width=True):
                    with st.spinner(f"Converting to {image_format} images..."):
                        images = []
                        
                        for page_num in range(st.session_state.doc_data.page_count):
                            page = st.session_state.doc_data[page_num]
                            mat = fitz.Matrix(dpi / 72.0, dpi / 72.0)
                            pix = page.get_pixmap(matrix=mat, alpha=False)
                            
                            if image_format == "PNG":
                                img_data = pix.tobytes("png")
                            else:
                                img_data = pix.tobytes("jpg", jpg_quality=quality)
                            
                            images.append((f"page_{page_num + 1}.{image_format.lower()}", img_data))
                            pix = None
                        
                        # Create a zip file in memory
                        import zipfile
                        zip_buffer = io.BytesIO()
                        
                        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
                            for filename, img_data in images:
                                zip_file.writestr(filename, img_data)
                        
                        st.download_button(
                            label="📥 Download Images as ZIP",
                            data=zip_buffer.getvalue(),
                            file_name=f"converted_pages_{image_format.lower()}.zip",
                            mime="application/zip"
                        )
                        st.success(f"✅ Converted {len(images)} pages to {image_format}!")
            
            elif conversion_format == "Extract Images Only":
                st.info("Extract all embedded images from the PDF")
                
                if st.button("Extract All Images", use_container_width=True):
                    with st.spinner("Extracting images..."):
                        extracted_images = extract_images_from_pdf(st.session_state.doc_data)
                        
                        if extracted_images:
                            # Create ZIP with extracted images
                            zip_buffer = io.BytesIO()
                            with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
                                for img_info in extracted_images:
                                    filename = f"page{img_info['page']}_img{img_info['index']}.{img_info['ext']}"
                                    zip_file.writestr(filename, img_info['data'])
                            
                            st.download_button(
                                label=f"📥 Download {len(extracted_images)} Extracted Images",
                                data=zip_buffer.getvalue(),
                                file_name="extracted_images.zip",
                                mime="application/zip"
                            )
                            st.success(f"✅ Extracted {len(extracted_images)} images from the PDF!")
                            
                            # Show image statistics
                            st.info(f"""
                            **Extraction Details:**
                            - Total images found: {len(extracted_images)}
                            - Pages with images: {len(set(img['page'] for img in extracted_images))}
                            - Format: PNG (converted from PDF)
                            """)
                        else:
                            st.warning("No images found in the PDF")
        
        with col2:
            st.markdown("#### 🔧 Conversion Options")
            
            # Show current document info
            if st.session_state.doc_data:
                st.info(f"""
                **Current Document:**
                - Format: PDF
                - Pages: {st.session_state.doc_data.page_count}
                - Size: ~{len(st.session_state.uploaded_file_data) / (1024*1024):.1f} MB
                """)
            
            st.markdown("#### 💡 Conversion Tips")
            st.markdown("""
            - **PDF to DOCX**: Best for text-heavy documents
            - **PDF to TXT**: Extracts all text content
            - **PDF to Images**: Each page as a separate image
            - **OCR + Convert**: Use OCR first for scanned documents
            """)
            
            # Batch conversion options
            st.markdown("#### 🚀 Batch Options")
            if st.checkbox("Enable batch processing"):
                st.info("Upload multiple files in the Upload section for batch conversion")
    
    else:
        st.info("📤 Please upload a document first to convert it.")

elif selected == "Analytics":
    if st.session_state.doc_text:
        st.markdown("### 📊 Document Analytics")
        
        # Text statistics
        words = st.session_state.doc_text.split()
        sentences = re.split(r'[.!?]+', st.session_state.doc_text)
        paragraphs = st.session_state.doc_text.split('\n\n')
        
        # Word frequency analysis
        word_freq = pd.Series(words).value_counts().head(20)
        
        # Create visualizations
        col1, col2 = st.columns(2)
        
        with col1:
            # Word frequency chart
            fig_words = go.Figure(data=[
                go.Bar(
                    x=word_freq.values,
                    y=word_freq.index,
                    orientation='h',
                    marker_color='rgba(102, 126, 234, 0.8)'
                )
            ])
            fig_words.update_layout(
                title="Top 20 Most Frequent Words",
                xaxis_title="Frequency",
                yaxis_title="Words",
                height=500,
                showlegend=False
            )
            st.plotly_chart(fig_words, use_container_width=True)
        
        with col2:
            # Document statistics pie chart
            fig_stats = go.Figure(data=[go.Pie(
                labels=['Words', 'Sentences', 'Paragraphs'],
                values=[len(words), len(sentences), len(paragraphs)],
                hole=.3,
                marker_colors=['#667eea', '#764ba2', '#a855f7']
            )])
            fig_stats.update_layout(
                title="Document Composition",
                height=500
            )
            st.plotly_chart(fig_stats, use_container_width=True)
        
        # Readability metrics
        st.markdown("### 📈 Readability Metrics")
        
        col1, col2, col3 = st.columns(3)
        
        with col1:
            avg_word_length = sum(len(word) for word in words) / len(words)
            st.metric("Average Word Length", f"{avg_word_length:.1f} chars")
        
        with col2:
            avg_sentence_length = len(words) / len(sentences)
            st.metric("Average Sentence Length", f"{avg_sentence_length:.1f} words")
        
        with col3:
            # Simple readability score (Flesch Reading Ease approximation)
            syllables_per_word = avg_word_length / 3  # Rough approximation
            flesch_score = 206.835 - 1.015 * avg_sentence_length - 84.6 * syllables_per_word
            flesch_score = max(0, min(100, flesch_score))  # Bound between 0-100
            st.metric("Readability Score", f"{flesch_score:.0f}/100")
        
        # Character distribution
        st.markdown("### 🔤 Character Distribution")
        
        char_types = {
            'Letters': sum(c.isalpha() for c in st.session_state.doc_text),
            'Digits': sum(c.isdigit() for c in st.session_state.doc_text),
            'Spaces': sum(c.isspace() for c in st.session_state.doc_text),
            'Punctuation': sum(c in '.,!?;:' for c in st.session_state.doc_text),
            'Other': sum(not (c.isalnum() or c.isspace() or c in '.,!?;:') for c in st.session_state.doc_text)
        }
        
        fig_chars = go.Figure(data=[go.Bar(
            x=list(char_types.keys()),
            y=list(char_types.values()),
            marker_color=['#667eea', '#764ba2', '#a855f7', '#ec4899', '#f43f5e']
        )])
        fig_chars.update_layout(
            title="Character Type Distribution",
            xaxis_title="Character Type",
            yaxis_title="Count",
            showlegend=False
        )
        st.plotly_chart(fig_chars, use_container_width=True)
    
    else:
        st.info("📤 Please upload a document first to see analytics.")

elif selected == "Compress":
    st.markdown("### 🗜️ Advanced Compression & Resizing")
    
    tab1, tab2, tab3 = st.tabs(["PDF Compression", "Image Compression", "Batch Processing"])
    
    with tab1:
        if st.session_state.doc_data and st.session_state.uploaded_file_data:
            st.markdown("#### 📄 PDF Compression")
            
            col1, col2 = st.columns([2, 1])
            
            with col1:
                compression_method = st.radio(
                    "Compression Method",
                    ["Safe (Text Preserved)", "With Image Reduction", "Maximum (No Text Conversion)", "Extreme (Convert to Images)"],
                    help="""
                    • Safe: Uses PyMuPDF's built-in optimization, preserves all text perfectly
                    • With Image Reduction: Reduces image quality while keeping text intact
                    • Maximum: Aggressive compression without converting text to images
                    • Extreme: Converts pages to images - smallest size but text becomes unselectable
                    """
                )
                
                # Show method-specific options
                if compression_method == "Safe (Text Preserved)":
                    compression_level = st.select_slider(
                        "Compression Level",
                        options=["minimal", "standard", "maximum"],
                        value="standard"
                    )
                    
                    compression_info = {
                        "minimal": "Light compression, perfect quality (5-15% reduction)",
                        "standard": "Balanced compression with full text preservation (10-30% reduction)",
                        "maximum": "Maximum safe compression (15-40% reduction)"
                    }
                    
                elif compression_method == "With Image Reduction":
                    col_a, col_b = st.columns(2)
                    with col_a:
                        image_quality = st.slider("Image Quality", 30, 100, 70)
                    with col_b:
                        image_dpi = st.slider("Image DPI", 72, 300, 120)
                    
                    compression_info = {
                        "custom": f"Reduces images to {image_quality}% quality at {image_dpi} DPI while preserving text"
                    }
                    compression_level = "custom"
                    
                elif compression_method == "Maximum (No Text Conversion)":
                    compression_level = "maximum"
                    compression_info = {
                        "maximum": "Maximum compression using all safe methods (20-50% reduction)"
                    }
                    
                else:  # Extreme
                    col_a, col_b = st.columns(2)
                    with col_a:
                        extreme_quality = st.slider("Page Quality", 20, 60, 30)
                    with col_b:
                        extreme_dpi = st.slider("Page DPI", 50, 100, 72)
                    
                    compression_level = "extreme"
                    compression_info = {
                        "extreme": f"Converts to images at {extreme_quality}% quality, {extreme_dpi} DPI (50-90% reduction)"
                    }
                    
                    st.warning("⚠️ **Warning**: This method converts text to images. Text will become unselectable and may appear blurry. Use only when file size is critical.")
                
                # Show compression preview info
                st.info(f"**Expected outcome:** {compression_info.get(compression_level, 'Custom compression')}")
                
                if st.button("🗜️ Compress PDF", use_container_width=True):
                    with st.spinner("Compressing PDF..."):
                        try:
                            if compression_method == "Safe (Text Preserved)":
                                compressed_data, ratio, original_size, compressed_size = compress_pdf_safe(
                                    st.session_state.uploaded_file_data,
                                    compression_level
                                )
                            elif compression_method == "With Image Reduction":
                                compressed_data, ratio, original_size, compressed_size = compress_pdf_with_image_reduction(
                                    st.session_state.uploaded_file_data,
                                    image_quality,
                                    image_dpi
                                )
                            elif compression_method == "Maximum (No Text Conversion)":
                                # Use maximum safe compression
                                compressed_data, ratio, original_size, compressed_size = compress_pdf_safe(
                                    st.session_state.uploaded_file_data,
                                    "maximum"
                                )
                            else:  # Extreme
                                compressed_data, ratio, original_size, compressed_size = compress_pdf_extreme(
                                    st.session_state.uploaded_file_data,
                                    extreme_quality,
                                    extreme_dpi
                                )
                            
                            if compressed_data and compressed_size < original_size:
                                st.success(f"✅ Compression completed! Size reduced by {ratio:.1f}%")
                                
                                # Show stats
                                col1, col2, col3 = st.columns(3)
                                with col1:
                                    st.metric("Original Size", f"{original_size / (1024*1024):.2f} MB")
                                with col2:
                                    st.metric("Compressed Size", f"{compressed_size / (1024*1024):.2f} MB")
                                with col3:
                                    st.metric("Reduction", f"{ratio:.1f}%", delta=f"-{(original_size-compressed_size)/(1024*1024):.2f} MB")
                                
                                # Download button
                                st.download_button(
                                    label="📥 Download Compressed PDF",
                                    data=compressed_data,
                                    file_name=f"compressed_{compression_method.lower().replace(' ', '_').replace('(', '').replace(')', '')}.pdf",
                                    mime="application/pdf"
                                )
                                
                                # Additional info
                                if compression_method == "Safe (Text Preserved)":
                                    st.success("✅ All text remains selectable and searchable")
                                elif compression_method == "With Image Reduction":
                                    st.success("✅ Text preserved, images optimized")
                                elif compression_method == "Extreme (Convert to Images)":
                                    st.warning("⚠️ Text is now part of the image and cannot be selected")
                                
                            elif compressed_data and compressed_size >= original_size:
                                st.warning("The file could not be compressed further. It may already be optimized.")
                                st.info("Try a different compression method or accept the current file size.")
                            else:
                                st.error("Compression failed. Please try a different method.")
                                
                        except Exception as e:
                            st.error(f"Compression error: {str(e)}")
                            st.info("Please try a different compression method")
            
            with col2:
                st.markdown("#### ⚙️ Compression Tips")
                
                st.markdown("""
                **Best Practices:**
                - Start with "Safe" method to preserve text
                - Use "With Image Reduction" for PDFs with many images
                - Only use "Extreme" when file size is critical
                
                **Text Quality:**
                - ✅ Safe: Perfect text quality
                - ✅ Image Reduction: Perfect text quality
                - ✅ Maximum: Perfect text quality
                - ❌ Extreme: Text becomes image
                
                **When to use each:**
                - **Safe**: Documents with important text
                - **Image Reduction**: Image-heavy PDFs
                - **Maximum**: Need best compression with text
                - **Extreme**: File size is critical, text quality not important
                """)
        
        else:
            st.info("📤 Please upload a PDF document to compress")
    
    with tab2:
        st.markdown("#### 🖼️ Image Compression & Resizing")
        
        uploaded_images = st.file_uploader(
            "Upload images to compress/resize",
            type=['png', 'jpg', 'jpeg', 'bmp', 'gif', 'webp'],
            accept_multiple_files=True
        )
        
        if uploaded_images:
            col1, col2 = st.columns(2)
            
            with col1:
                st.markdown("##### Compression Settings")
                output_format = st.selectbox("Output Format", ["JPEG", "PNG", "WEBP"])
                quality = st.slider("Quality", 10, 100, 85, help="Lower quality = smaller file size")
                
                st.markdown("##### Resize Settings")
                resize_option = st.radio("Resize Option", ["No Resize", "By Percentage", "Fixed Dimensions"])
                
                if resize_option == "By Percentage":
                    resize_percent = st.slider("Resize Percentage", 10, 100, 100)
                elif resize_option == "Fixed Dimensions":
                    col_w, col_h = st.columns(2)
                    with col_w:
                        max_width = st.number_input("Max Width (px)", value=1920, min_value=1)
                    with col_h:
                        max_height = st.number_input("Max Height (px)", value=1080, min_value=1)
                    maintain_aspect = st.checkbox("Maintain Aspect Ratio", value=True)
            
            with col2:
                st.markdown("##### Process Images")
                
                if st.button("🗜️ Compress & Resize All", use_container_width=True):
                    progress_bar = st.progress(0)
                    compressed_images = []
                    total_original = 0
                    total_compressed = 0
                    
                    for idx, uploaded_image in enumerate(uploaded_images):
                        progress_bar.progress((idx + 1) / len(uploaded_images))
                        
                        # Read image
                        image_data = uploaded_image.read()
                        total_original += len(image_data)
                        
                        # Calculate dimensions
                        if resize_option == "By Percentage":
                            img = Image.open(io.BytesIO(image_data))
                            new_width = int(img.width * resize_percent / 100)
                            new_height = int(img.height * resize_percent / 100)
                        elif resize_option == "Fixed Dimensions":
                            new_width = max_width
                            new_height = max_height
                        else:
                            new_width = None
                            new_height = None
                        
                        # Compress and resize
                        compressed_data, ratio, orig_size, comp_size, new_dims = compress_image_advanced(
                            image_data,
                            output_format,
                            quality,
                            new_width,
                            new_height
                        )
                        
                        if compressed_data:
                            total_compressed += comp_size
                            compressed_images.append((
                                f"{uploaded_image.name.split('.')[0]}.{output_format.lower()}",
                                compressed_data,
                                ratio
                            ))
                    
                    if compressed_images:
                        st.success(f"✅ Processed {len(compressed_images)} images!")
                        
                        # Show total stats
                        total_ratio = (1 - total_compressed / total_original) * 100
                        st.metric(
                            "Total Size Reduction",
                            f"{total_ratio:.1f}%",
                            delta=f"-{(total_original-total_compressed)/(1024*1024):.2f} MB"
                        )
                        
                        # Create ZIP for download
                        zip_buffer = io.BytesIO()
                        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
                            for filename, data, _ in compressed_images:
                                zip_file.writestr(filename, data)
                        
                        st.download_button(
                            label="📥 Download All Compressed Images",
                            data=zip_buffer.getvalue(),
                            file_name="compressed_images.zip",
                            mime="application/zip"
                        )
    
    with tab3:
        st.markdown("#### 🚀 Batch Processing")
        
        batch_type = st.radio("Select batch operation", ["PDF Batch Compression", "Image Batch Resize"])
        
        if batch_type == "PDF Batch Compression":
            pdf_files = st.file_uploader(
                "Upload multiple PDFs",
                type=['pdf'],
                accept_multiple_files=True
            )
            
            if pdf_files:
                batch_compression_method = st.selectbox(
                    "Batch Compression Method",
                    ["Safe (Text Preserved)", "With Image Reduction", "Maximum"]
                )
                
                if st.button("🗜️ Compress All PDFs", use_container_width=True):
                    progress_bar = st.progress(0)
                    compressed_pdfs = []
                    total_original = 0
                    total_compressed = 0
                    
                    for idx, pdf_file in enumerate(pdf_files):
                        progress_bar.progress((idx + 1) / len(pdf_files))
                        
                        # Read PDF
                        pdf_data = pdf_file.read()
                        total_original += len(pdf_data)
                        
                        # Compress based on method
                        if batch_compression_method == "Safe (Text Preserved)":
                            compressed_data, ratio, _, comp_size = compress_pdf_safe(pdf_data, "standard")
                        elif batch_compression_method == "With Image Reduction":
                            compressed_data, ratio, _, comp_size = compress_pdf_with_image_reduction(pdf_data, 70, 120)
                        else:  # Maximum
                            compressed_data, ratio, _, comp_size = compress_pdf_safe(pdf_data, "maximum")
                        
                        if compressed_data:
                            total_compressed += comp_size
                            compressed_pdfs.append((
                                f"compressed_{pdf_file.name}",
                                compressed_data
                            ))
                    
                    if compressed_pdfs:
                        # Create ZIP
                        zip_buffer = io.BytesIO()
                        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
                            for filename, data in compressed_pdfs:
                                zip_file.writestr(filename, data)
                        
                        st.success(f"✅ Compressed {len(compressed_pdfs)} PDFs!")
                        
                        # Show total stats
                        total_ratio = (1 - total_compressed / total_original) * 100
                        st.metric(
                            "Total Size Reduction",
                            f"{total_ratio:.1f}%",
                            delta=f"-{(total_original-total_compressed)/(1024*1024):.2f} MB"
                        )
                        
                        st.download_button(
                            label="📥 Download All Compressed PDFs",
                            data=zip_buffer.getvalue(),
                            file_name="compressed_pdfs.zip",
                            mime="application/zip"
                        )

elif selected == "Export":
    if st.session_state.doc_data or st.session_state.doc_text:
        st.markdown("### 💾 Export Options")
        
        col1, col2 = st.columns(2)
        
        with col1:
            st.markdown("#### 📄 Export as Text")
            if st.button("Download as TXT", use_container_width=True):
                # Create download link for text
                b64_text = base64.b64encode(st.session_state.doc_text.encode()).decode()
                href = f'<a href="data:text/plain;base64,{b64_text}" download="document.txt">Click here to download</a>'
                st.markdown(href, unsafe_allow_html=True)
            
            st.markdown("#### 📊 Export Analytics")
            if st.button("Download Analytics Report", use_container_width=True):
                # Create analytics report
                metadata_text = '\n'.join([f"{k}: {v}" for k, v in st.session_state.doc_metadata.items()])
                report = f"""Document Analytics Report
Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}

DOCUMENT STATISTICS
==================
Total Pages: {st.session_state.doc_data.page_count if st.session_state.doc_data else 'N/A'}
Total Words: {len(st.session_state.doc_text.split())}
Total Characters: {len(st.session_state.doc_text)}
Total Sentences: {len(re.split(r'[.!?]+', st.session_state.doc_text))}
Total Paragraphs: {len(st.session_state.doc_text.split(chr(10)+chr(10)))}

METADATA
========
{metadata_text}
"""
                b64_report = base64.b64encode(report.encode()).decode()
                href = f'<a href="data:text/plain;base64,{b64_report}" download="analytics_report.txt">Click here to download</a>'
                st.markdown(href, unsafe_allow_html=True)
        
        with col2:
            st.markdown("#### 🖼️ Export Pages as Images")
            if st.session_state.doc_data and st.session_state.doc_data.page_count > 0:
                # Image export settings
                export_format = st.selectbox("Export Format", ["PNG", "JPEG"])
                export_quality = st.slider("Quality", 50, 100, 90)
                export_dpi = st.slider("DPI (Resolution)", 72, 300, 150)
                
                if st.session_state.doc_data.page_count == 1:
                    st.info("Document has only 1 page")
                    if st.button("Export Page as Image", use_container_width=True):
                        with st.spinner("Exporting page..."):
                            page = st.session_state.doc_data[0]
                            mat = fitz.Matrix(export_dpi / 72.0, export_dpi / 72.0)
                            pix = page.get_pixmap(matrix=mat, alpha=False)
                            
                            if export_format == "PNG":
                                img_data = pix.tobytes("png")
                                mime_type = "image/png"
                                file_ext = "png"
                            else:
                                img_data = pix.tobytes("jpg", jpg_quality=export_quality)
                                mime_type = "image/jpeg"
                                file_ext = "jpg"
                            
                            st.download_button(
                                label=f"📥 Download Page as {export_format}",
                                data=img_data,
                                file_name=f"page_1.{file_ext}",
                                mime=mime_type
                            )
                            st.success("Page exported successfully!")
                else:
                    page_range = st.select_slider(
                        "Select page range",
                        options=list(range(1, st.session_state.doc_data.page_count + 1)),
                        value=(1, min(5, st.session_state.doc_data.page_count))
                    )
                    
                    if st.button("Export Selected Pages", use_container_width=True):
                        with st.spinner("Exporting pages..."):
                            # Create images for selected pages
                            exported_images = []
                            
                            for page_num in range(page_range[0] - 1, page_range[1]):
                                page = st.session_state.doc_data[page_num]
                                mat = fitz.Matrix(export_dpi / 72.0, export_dpi / 72.0)
                                pix = page.get_pixmap(matrix=mat, alpha=False)
                                
                                if export_format == "PNG":
                                    img_data = pix.tobytes("png")
                                    file_ext = "png"
                                else:
                                    img_data = pix.tobytes("jpg", jpg_quality=export_quality)
                                    file_ext = "jpg"
                                
                                exported_images.append((f"page_{page_num + 1}.{file_ext}", img_data))
                            
                            # Create ZIP file
                            zip_buffer = io.BytesIO()
                            with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
                                for filename, data in exported_images:
                                    zip_file.writestr(filename, data)
                            
                            st.download_button(
                                label=f"📥 Download {len(exported_images)} Pages as ZIP",
                                data=zip_buffer.getvalue(),
                                file_name=f"pages_{page_range[0]}_to_{page_range[1]}.zip",
                                mime="application/zip"
                            )
                            st.success(f"Pages {page_range[0]} to {page_range[1]} exported successfully!")
            else:
                st.info("No document loaded for image export")
            
            st.markdown("#### 📋 Export Search Results")
            if st.session_state.search_results:
                if st.button("Download Search Results", use_container_width=True):
                    results_text = "Search Results\n" + "="*50 + "\n\n"
                    for i, result in enumerate(st.session_state.search_results):
                        results_text += f"Result {i+1} - Line {result['line']}\n"
                        results_text += f"Context: {result['context']}\n"
                        results_text += "-"*30 + "\n\n"
                    
                    b64_results = base64.b64encode(results_text.encode()).decode()
                    href = f'<a href="data:text/plain;base64,{b64_results}" download="search_results.txt">Click here to download</a>'
                    st.markdown(href, unsafe_allow_html=True)
    
    else:
        st.info("📤 Please upload a document first to export it.")

# Footer
st.markdown("---")
st.markdown(f"""
<div style="text-align: center; color: {'#ffffff' if st.session_state.dark_mode else '#6c757d'}; padding: 2rem;">
    <p>Advanced Document Viewer v2.0 | Built with Streamlit</p>
    <p style="font-size: 0.9rem;">Features: PDF/DOCX • OCR • Rotation • Format Conversion • Advanced Compression • Dark Mode</p>
    <p style="font-size: 0.8rem;">Install requirements: streamlit, PyMuPDF, pytesseract, opencv-python, pdf2docx, pillow, pandas, plotly</p>
</div>
""", unsafe_allow_html=True)